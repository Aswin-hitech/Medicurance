from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text):
    return doc.add_paragraph(text)

def generate_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('MediCurance: Project Overview & Architecture', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 1. Summary of Project
    add_heading(doc, '1. Summary of the Project')
    add_paragraph(doc, 
        "MediCurance is an AI-powered medical reimbursement validation system designed to automate and streamline healthcare claim processing for government employees and pensioners. "
        "The traditional process of claiming medical reimbursements is often tedious, manual, and prone to errors. MediCurance solves this by integrating Optical Character Recognition (OCR) "
        "and a Retrieval-Augmented Generation (RAG) AI pipeline. It allows users to upload their medical bills, automatically extracts the relevant information, checks the claims against "
        "government annexure rules, and provides a clear 'Eligible' or 'Not Eligible' recommendation for human officers to review. This ensures transparency, efficiency, and accuracy."
    )
    
    # 2. Tech Stacks
    add_heading(doc, '2. Technology Stack & Their Roles')
    stacks = [
        ("Flask (Python)", "Serves as the backend web framework handling routing, HTTP requests, and connecting the frontend with the AI/Database layers."),
        ("MongoDB Atlas", "A NoSQL database used to store user profiles, claim histories, hospital lists, and audit logs flexibly."),
        ("Supabase", "Cloud storage provider used to securely store uploaded medical bills and documents."),
        ("Groq API (LLM)", "The large language model (AI engine) used to reason over the extracted text and make eligibility decisions."),
        ("FAISS / sentence-transformers", "Used to create and query a local vector database. It turns government annexure rules into embeddings so the AI can quickly fetch the most relevant rules during a claim."),
        ("pdfplumber / pytesseract / OCR.Space", "Optical Character Recognition tools used to extract text and data from uploaded PDFs and image-based medical bills."),
        ("Jinja2 / HTML / CSS / JS", "The frontend templating engine and technologies used to build the user, officer, and admin dashboards."),
        ("bcrypt / PyJWT", "Used for security, specifically hashing passwords and managing secure session tokens.")
    ]
    for tech, role in stacks:
        p = doc.add_paragraph()
        p.add_run(f"{tech}: ").bold = True
        p.add_run(role)

    # 3. Workflow of the Project (Story Manner)
    add_heading(doc, '3. The Workflow (A Story)')
    add_paragraph(doc, 
        "Imagine a retired government teacher, Mr. Sharma. He recently underwent surgery and needs to claim his medical expenses. "
        "Instead of waiting in long lines at a government office, he logs into the MediCurance portal using his mobile number and an OTP. "
        "He fills out a simple form and uploads a photo of his hospital bill. "
    )
    add_paragraph(doc,
        "The moment he clicks 'Submit', the magic begins. MediCurance takes the image and passes it to its OCR engines, which carefully read the text, "
        "extracting the hospital name, the treatment details, and the total amount. But reading the bill isn't enough; the system needs to know if the treatment is covered. "
        "So, the AI searches its internal 'brain'—a vector database filled with government healthcare rules (Annexures)—to find the exact policies related to Mr. Sharma's surgery."
    )
    add_paragraph(doc,
        "With the bill details and the official rules in hand, the AI (powered by Groq) acts as a virtual auditor. It compares the two and concludes, "
        "'Yes, this hospital is networked, and this surgery is covered.' It generates a Confidence Score of 95% and marks the claim as 'Eligible.' "
    )
    add_paragraph(doc,
        "However, AI doesn't make the final call—humans do. An appointed Government Officer logs into their dashboard and sees Mr. Sharma's claim. "
        "Instead of manually verifying the rules, the officer simply reviews the AI's summary and reasoning. Seeing the AI's approval, the officer clicks 'Approve', "
        "and Mr. Sharma is immediately notified that his reimbursement is on its way. What used to take months now takes minutes!"
    )

    # 4. Architecture & Modules
    add_heading(doc, '4. Architecture, Modules & Files')
    add_paragraph(doc, "The project follows a modular Monolith architecture using Flask Blueprints:")
    modules = [
        ("app.py", "The main entry point of the application. It initializes the Flask server, registers blueprints, and runs startup health checks."),
        ("config/settings.py", "Centralized configuration management. Loads environment variables (like API keys and DB URIs) safely."),
        ("blueprints/", "Contains the routing logic separated by user roles. `auth.py` handles login/registration. `user.py` handles user dashboards and claim submissions. `officer.py` handles the review dashboard. `admin.py` is for managing the system."),
        ("services/", "The core business logic layer. Files here do the heavy lifting:\n"
                      "- `ocr_service.py`: Handles text extraction from files.\n"
                      "- `rag_service.py`: Handles querying the vector DB and prompting the AI.\n"
                      "- `otp_service.py`: Generates and verifies OTPs for login.\n"
                      "- `storage_service.py`: Uploads files to Supabase.\n"
                      "- `claim_processing_service.py`: The orchestrator that ties OCR, RAG, and Database together when a claim is submitted."),
        ("database/", "Contains repositories for database interactions (e.g., `mongo_client.py` sets up MongoDB connections, `user_repository.py` handles user queries)."),
        ("vectorstore/ & build_vector.py", "The `build_vector.py` script reads PDFs from `resources/annexures/`, converts their text into embeddings, and saves them in the `vectorstore/` directory for fast AI retrieval."),
        ("utils/", "Helper scripts like `logger.py` for audit logging and `rate_limiter.py` to prevent spam attacks.")
    ]
    for module, desc in modules:
        p = doc.add_paragraph()
        p.add_run(f"{module}: ").bold = True
        p.add_run(desc)
        
    add_paragraph(doc, "Note on Agents: Currently, the system uses a standard RAG pipeline (Retrieval-Augmented Generation) rather than autonomous AI agents. The AI performs structured zero-shot analyses rather than looping and planning on its own. A 'Multi-agent AI system' is planned for the future.")

    # 5. What happens in each step
    add_heading(doc, '5. Step-by-Step Execution Flow')
    add_paragraph(doc, "When a user submits a claim, the following exact steps occur:")
    steps = [
        "1. Request Received: The frontend sends a POST request with form data and the uploaded file to the `/submit_claim` endpoint in `blueprints/user.py`.",
        "2. File Upload: The `storage_service.py` takes the file and uploads it securely to Supabase, returning a public URL.",
        "3. Data Extraction (OCR): `ocr_service.py` downloads the file (or uses the stream) and extracts all readable text from the bill.",
        "4. Rule Retrieval (RAG): `rag_service.py` takes the extracted text and queries the local FAISS vector database to find the most relevant government annexure rules.",
        "5. AI Validation: `rag_service.py` sends a combined prompt (containing the bill text and the retrieved rules) to the Groq LLM. The LLM returns a structured JSON response (Eligible/Not Eligible, Score, Reason).",
        "6. Database Storage: The claim details, file URL, and AI validation results are saved into the MongoDB `claims` collection.",
        "7. Review Phase: An officer logs into the `officer` blueprint, fetches the claim from MongoDB, reads the AI's reasoning, and clicks Approve/Reject.",
        "8. Notification: The system updates the claim status in the database and the user can see the final decision on their dashboard."
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Bullet')

    # Save
    doc.save('MediCurance_Project_Report.docx')

if __name__ == "__main__":
    generate_report()
    print("Report generated successfully at MediCurance_Project_Report.docx")
