import logging
import os
import shutil
import tempfile
from pathlib import Path

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.opc.exceptions import PackageNotFoundError
    from docx.shared import Inches, Pt
except Exception:  # pragma: no cover - optional dependency fallback
    Document = None
    WD_ALIGN_PARAGRAPH = None
    Inches = None
    Pt = None

    class PackageNotFoundError(Exception):
        pass


def _normalize_value(value):
    if value is None:
        return ""
    return str(value)


def _replace_text(text, data):
    updated = text
    for key, value in data.items():
        updated = updated.replace(f"{{{{{key}}}}}", _normalize_value(value))
    return updated


def _replace_in_paragraph(paragraph, data):
    if not paragraph.text:
        return

    original_text = paragraph.text
    run_updated = False

    for run in paragraph.runs:
        if run.text:
            new_text = _replace_text(run.text, data)
            if new_text != run.text:
                run.text = new_text
                run_updated = True

    if any(f"{{{{{key}}}}}" in paragraph.text for key in data):
        paragraph.text = _replace_text(original_text, data)
    elif not run_updated:
        return


def _replace_in_cell(cell, data):
    for paragraph in cell.paragraphs:
        _replace_in_paragraph(paragraph, data)
    for table in cell.tables:
        _replace_in_table(table, data)


def _replace_in_table(table, data):
    for row in table.rows:
        for cell in row.cells:
            _replace_in_cell(cell, data)


def _replace_in_container(container, data):
    for paragraph in container.paragraphs:
        _replace_in_paragraph(paragraph, data)
    for table in container.tables:
        _replace_in_table(table, data)


def _replace_in_document(doc, data):
    _replace_in_container(doc, data)

    for section in doc.sections:
        _replace_in_container(section.header, data)
        _replace_in_container(section.footer, data)


def _apply_paragraph_style(paragraph):
    if not paragraph.runs:
        return
    for run in paragraph.runs:
        font = run.font
        font.name = "Segoe UI"
        font.size = Pt(11) if Pt else font.size


def _append_letter_section(doc, data):
    body = str(data.get("letter_body") or "").strip()
    if not body:
        return

    for index, raw_line in enumerate(body.splitlines()):
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if WD_ALIGN_PARAGRAPH else paragraph.alignment
        run = paragraph.add_run(line)
        if Pt:
            run.font.name = "Segoe UI"
            run.font.size = Pt(11)
        if index == 0 and line.lower().startswith("date:"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if WD_ALIGN_PARAGRAPH else paragraph.alignment


def _append_meta_block(doc, data):
    meta_rows = [
        ("Date", data.get("current_date")),
        ("Reference Number", data.get("letter_reference") or data.get("claim_reference")),
        ("Claim ID", data.get("claim_id")),
        ("Beneficiary", data.get("beneficiary_name")),
        ("PPO Number", data.get("ppo_number") or data.get("employee_id")),
        ("Hospital", data.get("hospital")),
        ("Amount", data.get("amount")),
        ("Status", data.get("status")),
        ("Officer", data.get("officer_name")),
        ("Designation", data.get("officer_designation")),
    ]

    for label, value in meta_rows:
        if not value:
            continue
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if WD_ALIGN_PARAGRAPH else paragraph.alignment
        run = paragraph.add_run(f"{label}: {value}")
        if Pt:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)


def _append_stamp(doc, data):
    if WD_ALIGN_PARAGRAPH is None:
        return

    stamp_path = data.get("stamp_path")
    if not stamp_path:
        return

    stamp_file = Path(stamp_path)
    if not stamp_file.exists():
        return

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    try:
        run.add_picture(str(stamp_file), width=Inches(1.35))
    except Exception as exc:
        logger.warning("Stamp insertion skipped: %s", exc)


def generate_letter_from_template(template_path, data):
    """
    Load a DOCX template, replace placeholders, and save a populated DOCX file.
    Returns the generated DOCX path.
    """
    if Document is None:
        raise ImportError("python-docx is not installed")

    template_path = Path(template_path)
    if not template_path.exists() or template_path.stat().st_size == 0:
        raise FileNotFoundError(f"Template missing or empty: {template_path}")

    try:
        doc = Document(str(template_path))
    except PackageNotFoundError as exc:
        raise FileNotFoundError(f"Invalid DOCX template: {template_path}") from exc

    payload = data or {}
    _replace_in_document(doc, payload)
    _append_meta_block(doc, payload)
    _append_letter_section(doc, payload)
    _append_stamp(doc, payload)

    output = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    output_path = output.name
    output.close()
    doc.save(output_path)
    return output_path


def convert_docx_to_pdf(docx_path):
    """
    Convert a DOCX file to PDF.
    Tries docx2pdf first, then LibreOffice/soffice, then raises.
    Returns the generated PDF path.
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX not found: {docx_path}")

    output_dir = Path(tempfile.mkdtemp(prefix="medicurance_letter_pdf_"))
    pdf_path = output_dir / f"{docx_path.stem}.pdf"

    try:
        try:
            from docx2pdf import convert

            convert(str(docx_path), str(output_dir))
        except Exception as first_exc:
            soffice = shutil.which("soffice") or shutil.which("libreoffice")
            if not soffice:
                raise first_exc

            import subprocess

            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf:writer_pdf_Export", "--outdir", str(output_dir), str(docx_path)],
                capture_output=True,
                text=True,
                check=False,
            )
            if result.returncode != 0:
                raise RuntimeError(
                    f"LibreOffice conversion failed: {result.stderr.strip() or result.stdout.strip()}"
                ) from first_exc

        if not pdf_path.exists():
            raise FileNotFoundError(f"PDF conversion did not produce output: {pdf_path}")

        return str(pdf_path)
    except Exception:
        shutil.rmtree(output_dir, ignore_errors=True)
        raise
    finally:
        try:
            os.remove(docx_path)
        except Exception:
            pass
