import os
import logging
import tempfile
import requests
from datetime import datetime
from pathlib import Path
from flask import current_app

# ReportLab Imports
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# Python-Docx Imports
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Storage Imports
from services.storage_service import upload_file as upload_to_storage
from config.settings import Config

logger = logging.getLogger(__name__)

def _download_image(url: str) -> str | None:
    """Download an image from public URL to a temp file for ReportLab/Docx rendering."""
    if not url or not url.startswith("http"):
        return None
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        suffix = Path(url.split("?")[0]).suffix or ".png"
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        temp_file.write(response.content)
        temp_file.close()
        return temp_file.name
    except Exception as e:
        logger.warning(f"[DocGen] Failed to download image from {url}: {e}")
        return None

def _make_digit_boxes(value_str, box_count=12):
    val = "".join(ch for ch in str(value_str or "") if ch.isalnum())
    val = val.ljust(box_count)[:box_count]
    
    # Cells containing single character paragraphs
    cell_style = ParagraphStyle('DigCell', fontName='Helvetica-Bold', fontSize=10, leading=11, alignment=1)
    data = [[Paragraph(f"<b>{ch}</b>" if ch != " " else "", cell_style) for ch in val]]
    col_widths = [14] * box_count
    t = Table(data, colWidths=col_widths, rowHeights=[16])
    t.setStyle(TableStyle([
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#000000')),
        ('INNERGRID', (0,0), (-1,-1), 1, colors.HexColor('#000000')),
        ('BOTTOMPADDING', (0,0), (-1,-1), 1),
        ('TOPPADDING', (0,0), (-1,-1), 1),
        ('LEFTPADDING', (0,0), (-1,-1), 0),
        ('RIGHTPADDING', (0,0), (-1,-1), 0),
    ]))
    return t

def _format_dob_digits(dob_str):
    if not dob_str:
        return ""
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            dt = datetime.strptime(str(dob_str).split("T")[0], fmt)
            return dt.strftime("%d%m%Y")
        except Exception:
            pass
    return "".join(ch for ch in str(dob_str) if ch.isdigit())

def generate_pdf_application(claim_data: dict, photo_local_path: str | None, output_path: str) -> bool:
    """Generate the official government application PDF using ReportLab with exact user-requested layout."""
    try:
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            leftMargin=36,
            rightMargin=36,
            topMargin=36,
            bottomMargin=36
        )
        
        styles = getSampleStyleSheet()
        
        # Styles definition
        title_style = ParagraphStyle(
            'GovTitle',
            fontName='Helvetica-Bold',
            fontSize=16,
            leading=20,
            alignment=1,
            textColor=colors.HexColor('#12284c')
        )
        label_style = ParagraphStyle(
            'GovLabel',
            fontName='Helvetica-Bold',
            fontSize=9,
            leading=12,
            textColor=colors.HexColor('#12284c')
        )
        value_style = ParagraphStyle(
            'GovValue',
            fontName='Helvetica',
            fontSize=9.5,
            leading=12,
            textColor=colors.HexColor('#000000')
        )
        dec_style = ParagraphStyle(
            'GovDec',
            fontName='Helvetica-Oblique',
            fontSize=9,
            leading=14,
            textColor=colors.HexColor('#333333')
        )
        
        story = []
        
        # 1. Header Row (Left Logo & Text, Right Photo Box)
        # Check if website logo exists
        logo_path = Path(__file__).resolve().parent.parent / "static" / "assets" / "medicurance_logo.png"
        logo_flowable = None
        if logo_path.exists():
            try:
                logo_flowable = Image(str(logo_path), width=48, height=48)
            except Exception:
                pass
                
        # Header text and Registration (PPO) boxes
        ppo_num = claim_data.get("ppo_number") or ""
        ppo_boxes = _make_digit_boxes(ppo_num, box_count=10)
        
        sub_date = datetime.now().strftime("%d / %m / %Y")
        
        header_text_table_data = [
            [Paragraph("<b>GOVERNMENT OF TAMIL NADU</b>", label_style)],
            [Paragraph("<font size=8.5 color='#666666'>PENSIONERS HEALTH SCHEME 2026</font>", label_style)],
            [Spacer(1, 4)],
            [Table([
                [Paragraph("<b>Reg. No. (PPO):</b>", label_style), ppo_boxes]
            ], colWidths=[90, 150], style=[
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('BOTTOMPADDING', (0,0), (-1,-1), 0),
            ])],
            [Spacer(1, 4)],
            [Paragraph(f"<b>Date of Issue:</b> {sub_date}", label_style)]
        ]
        
        header_text_table = Table(header_text_table_data, colWidths=[350])
        header_text_table.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('LEFTPADDING', (0,0), (-1,-1), 0),
            ('BOTTOMPADDING', (0,0), (-1,-1), 1),
            ('TOPPADDING', (0,0), (-1,-1), 1),
        ]))
        
        # Merge logo + header text
        if logo_flowable:
            left_header = Table([[logo_flowable, header_text_table]], colWidths=[60, 360])
            left_header.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
            ]))
        else:
            left_header = header_text_table
            
        # Photo Frame (Right side)
        photo_flowable = None
        if photo_local_path and os.path.exists(photo_local_path):
            try:
                photo_flowable = Image(photo_local_path, width=75, height=90)
            except Exception:
                pass
                
        if not photo_flowable:
            photo_data = [[Paragraph("<font size=8 color='#888888'>Passport<br/>Photo<br/>Here</font>", label_style)]]
            photo_flowable = Table(photo_data, colWidths=[75], rowHeights=[90])
            photo_flowable.setStyle(TableStyle([
                ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#000000')),
                ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ]))
            
        header_master = Table([[left_header, photo_flowable]], colWidths=[440, 80])
        header_master.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('ALIGN', (1,0), (1,0), 'RIGHT'),
            ('LEFTPADDING', (0,0), (-1,-1), 0),
        ]))
        story.append(header_master)
        story.append(Spacer(1, 10))
        
        # 2. Main Title & Horizontal Divider
        story.append(Paragraph("<b>MEDICAL REIMBURSEMENT APPLICATION FORM</b>", title_style))
        story.append(Spacer(1, 6))
        
        # Line Divider
        line_table = Table([[""]], colWidths=[520], rowHeights=[2])
        line_table.setStyle(TableStyle([
            ('LINEBELOW', (0,0), (-1,-1), 1.5, colors.HexColor('#000000')),
            ('BOTTOMPADDING', (0,0), (-1,-1), 0),
            ('TOPPADDING', (0,0), (-1,-1), 0),
        ]))
        story.append(line_table)
        story.append(Spacer(1, 12))
        
        # 3. Personal Fields Layout
        fullname = claim_data.get("beneficiary_name") or ""
        # Split surname vs name
        name_parts = fullname.split(" ")
        surname_val = name_parts[-1] if len(name_parts) > 1 else ""
        name_val = " ".join(name_parts[:-1]) if len(name_parts) > 1 else fullname
        
        # Aadhaar digits
        aadhaar_num = claim_data.get("aadhaar_number") or ""
        aadhaar_boxes = _make_digit_boxes(aadhaar_num, box_count=12)
        
        # Date of birth digits
        dob_digits = _format_dob_digits(claim_data.get("dob"))
        dob_boxes = _make_digit_boxes(dob_digits, box_count=8)
        
        # Gender checkboxes
        gender_val = str(claim_data.get("gender") or "").strip().lower()
        male_box = "[✓]" if gender_val == "male" else "[  ]"
        female_box = "[✓]" if gender_val == "female" else "[  ]"
        gender_text = f"<b>Gender:</b>  {male_box} Male   {female_box} Female"
        
        # Render Fields
        fields_data = [
            [Paragraph("<b>Surname:</b>", label_style), Paragraph(f"<u>{surname_val}</u>", value_style)],
            [Paragraph("<b>Name:</b>", label_style), Paragraph(f"<u>{name_val}</u>", value_style)],
            [Paragraph("<b>Aadhar Card No.:</b>", label_style), aadhaar_boxes],
            [Paragraph("<b>Date of Birth:</b>", label_style), Table([[dob_boxes, Paragraph("<font size=8.5 color='#555555'>Format (DD/MM/YYYY)</font>", value_style)]], colWidths=[120, 150], style=[('VALIGN', (0,0), (-1,-1), 'MIDDLE'), ('LEFTPADDING', (0,0), (-1,-1), 0)])],
            [Paragraph(gender_text, label_style), Paragraph(f"<b>Phone:</b> <u>{claim_data.get('mobile_number') or ''}</u>", label_style)]
        ]
        
        fields_table = Table(fields_data, colWidths=[120, 400])
        fields_table.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
            ('TOPPADDING', (0,0), (-1,-1), 6),
            ('LEFTPADDING', (0,0), (-1,-1), 0),
        ]))
        story.append(fields_table)
        story.append(Spacer(1, 10))
        
        # 4. Supplementary Details Sections
        # Section Builder Helper
        def add_sub_section(sec_title, sec_data):
            section_label_style = ParagraphStyle(
                'SecTitle', fontName='Helvetica-Bold', fontSize=10, leading=12, textColor=colors.HexColor('#12284c'),
                spaceBefore=10, spaceAfter=4
            )
            story.append(Paragraph(f"<b>{sec_title}</b>", section_label_style))
            
            grid_data = []
            keys = list(sec_data.keys())
            for idx in range(0, len(keys), 2):
                k1 = keys[idx]
                v1 = sec_data[k1]
                p_k1 = Paragraph(str(k1), label_style)
                p_v1 = Paragraph(str(v1 or 'N/A'), value_style)
                if idx + 1 < len(keys):
                    k2 = keys[idx+1]
                    v2 = sec_data[k2]
                    p_k2 = Paragraph(str(k2), label_style)
                    p_v2 = Paragraph(str(v2 or 'N/A'), value_style)
                else:
                    p_k2, p_v2 = "", ""
                grid_data.append([p_k1, p_v1, p_k2, p_v2])
                
            grid_table = Table(grid_data, colWidths=[120, 140, 120, 140])
            grid_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('LINEBELOW', (0,0), (-1,-1), 0.5, colors.HexColor('#e5e5e5')),
                ('TOPPADDING', (0,0), (-1,-1), 4),
                ('BOTTOMPADDING', (0,0), (-1,-1), 4),
                ('LEFTPADDING', (0,0), (-1,-1), 4),
            ]))
            story.append(grid_table)
            
        # 1. Pensioner Details
        add_sub_section("I. Pensioner & Department Info", {
            "Retirement Date": claim_data.get("retirement_date"),
            "Category": claim_data.get("pension_category"),
            "Department": claim_data.get("department"),
            "Designation": claim_data.get("designation"),
            "Treasury Office": claim_data.get("treasury_office"),
            "Address": claim_data.get("address")
        })
        
        # 2. Treatment & Hospital Info
        add_sub_section("II. Treatment Details", {
            "Patient Name": claim_data.get("name"),
            "Relationship": claim_data.get("relationship"),
            "Disease Name": claim_data.get("disease"),
            "Diagnosis": claim_data.get("diagnosis"),
            "Procedure": claim_data.get("surgery_type"),
            "Hospital": claim_data.get("hospital"),
            "Admission Date": claim_data.get("admission_date"),
            "Discharge Date": claim_data.get("discharge_date"),
        })
        
        # 3. Claim & Bank Details
        add_sub_section("III. Reimbursement & Bank Info", {
            "Claim Type": claim_data.get("claim_type"),
            "Claim Amount": f"Rs. {claim_data.get('amount', '0.00')}",
            "Bank Name": claim_data.get("bank_name"),
            "Branch": claim_data.get("branch"),
            "Account Number": claim_data.get("account_number"),
            "IFSC Code": claim_data.get("ifsc"),
        })
        
        # 5. Undertaking Box (Declaration)
        story.append(Spacer(1, 15))
        dec_box_text = (
            "<b>UNDERTAKING</b><br/>"
            "I hereby declare that the information furnished above is true and correct to the best of "
            "my knowledge. I understand that furnishing false information may result in rejection "
            "of my claim and legal action as per applicable rules."
        )
        dec_box_data = [[Paragraph(dec_box_text, dec_style)]]
        dec_table = Table(dec_box_data, colWidths=[520])
        dec_table.setStyle(TableStyle([
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#000000')),
            ('TOPPADDING', (0,0), (-1,-1), 8),
            ('BOTTOMPADDING', (0,0), (-1,-1), 8),
            ('LEFTPADDING', (0,0), (-1,-1), 10),
            ('RIGHTPADDING', (0,0), (-1,-1), 10),
        ]))
        
        # Signature Line
        sig_name = claim_data.get("signature_name") or claim_data.get("beneficiary_name") or "Beneficiary Signature"
        sig_data = [[
            Paragraph(f"<b>Date:</b> {datetime.now().strftime('%d/%m/%Y')}", value_style),
            Paragraph(f"<b>Signature:</b> ___________________________<br/><font size=7 color='#555555'>({sig_name})</font>", ParagraphStyle('SigRight', fontName='Helvetica', alignment=2))
        ]]
        sig_table = Table(sig_data, colWidths=[200, 320])
        sig_table.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'BOTTOM'),
            ('TOPPADDING', (0,0), (-1,-1), 14),
            ('LEFTPADDING', (0,0), (-1,-1), 0),
        ]))
        
        story.append(KeepTogether([dec_table, sig_table]))
        
        # Build PDF document
        doc.build(story)
        return True
    except Exception as e:
        logger.error(f"[DocGen] ReportLab PDF generation failed: {e}", exc_info=True)
        return False

def generate_docx_application(claim_data: dict, photo_local_path: str | None, output_path: str) -> bool:
    """Generate the official government application DOCX using python-docx."""
    try:
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Style configurations
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(9.5)

        # Add Title
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.add_run("GOVERNMENT OF TAMIL NADU\nMEDICAL REIMBURSEMENT APPLICATION\n")
        run.bold = True
        run.font.size = Pt(13)
        
        run_sub = h.add_run("UNDER TN PENSIONERS HEALTH SCHEME 2026")
        run_sub.font.size = Pt(9.5)
        run_sub.font.color.rgb = docx.shared.RGBColor(100, 116, 139) if 'docx' in globals() else None

        # Embed photo if available
        if photo_local_path and os.path.exists(photo_local_path):
            try:
                p_table = doc.add_table(rows=1, cols=2)
                p_table.autofit = False
                p_table.columns[0].width = Inches(5.5)
                p_table.columns[1].width = Inches(1.5)
                
                # Title text
                p_table.cell(0, 0).text = f"Application Ref No: {claim_data.get('claim_number', 'N/A')}\nDate of Submission: {datetime.now().strftime('%d %B %Y')}"
                
                # Image
                p_cell = p_table.cell(0, 1)
                p_para = p_cell.paragraphs[0]
                p_run = p_para.add_run()
                p_run.add_picture(photo_local_path, width=Inches(1.0), height=Inches(1.0))
            except Exception as img_err:
                logger.warning(f"[DocGen] Failed to embed DOCX image: {img_err}")
                doc.add_paragraph(f"Application Ref No: {claim_data.get('claim_number', 'N/A')}")
        else:
            doc.add_paragraph(f"Application Ref No: {claim_data.get('claim_number', 'N/A')}")

        def add_docx_section(title, data_dict):
            p = doc.add_paragraph()
            r = p.add_run(title)
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = docx.shared.RGBColor(10, 91, 211) if 'docx' in globals() else None
            
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            for key, val in data_dict.items():
                row = table.add_row()
                row.cells[0].paragraphs[0].add_run(str(key)).bold = True
                row.cells[1].paragraphs[0].add_run(str(val or 'N/A'))

        # Add Details
        add_docx_section("1. Beneficiary Details", {
            "Beneficiary Full Name": claim_data.get("beneficiary_name"),
            "PPO Number": claim_data.get("ppo_number"),
            "Aadhaar Number": claim_data.get("aadhaar_number"),
            "Mobile": claim_data.get("mobile_number"),
            "Email": claim_data.get("email_address"),
            "Date of Birth": claim_data.get("dob"),
            "Retirement Date": claim_data.get("retirement_date"),
            "Pension Type": claim_data.get("pension_category"),
            "Department": claim_data.get("department"),
            "Designation": claim_data.get("designation"),
            "Treasury Office": claim_data.get("treasury_office"),
            "Address": claim_data.get("address")
        })

        add_docx_section("2. Treatment Details", {
            "Patient Name": claim_data.get("name"),
            "Relationship": claim_data.get("relationship"),
            "Disease": claim_data.get("disease"),
            "Diagnosis": claim_data.get("diagnosis"),
            "Procedure Type": claim_data.get("surgery_type"),
            "Treatment Category": claim_data.get("treatment_category"),
            "Admission Date": claim_data.get("admission_date"),
            "Discharge Date": claim_data.get("discharge_date"),
            "Treating Doctor": claim_data.get("doctor_name"),
            "Hospital": claim_data.get("hospital")
        })

        add_docx_section("3. Claim Details", {
            "Claim Type": claim_data.get("claim_type"),
            "Amount (INR)": f"Rs. {claim_data.get('amount', '0.00')}",
            "Previous Claims count": claim_data.get("prev_claim_count", "0"),
            "Prev Claim Reference": claim_data.get("prev_claim_ref"),
            "Emergency Case": claim_data.get("emergency_case", "No"),
            "Treatment Period": claim_data.get("treatment_period")
        })

        add_docx_section("4. Bank Details", {
            "Bank Name": claim_data.get("bank_name"),
            "Branch": claim_data.get("branch"),
            "Account Number": claim_data.get("account_number"),
            "IFSC": claim_data.get("ifsc"),
            "MICR": claim_data.get("micr"),
            "Account Type": claim_data.get("account_type"),
            "ECS Enabled": claim_data.get("ecs_enabled")
        })

        # Declaration
        doc.add_paragraph("\n5. Declaration")
        d_p = doc.add_paragraph()
        d_r = d_p.add_run("I hereby declare that the information furnished above is true and correct to the best of my knowledge. I understand that furnishing false information may result in rejection of my claim and legal action.")
        d_r.italic = True
        
        doc.add_paragraph(f"Date: {claim_data.get('signature_date') or datetime.now().strftime('%Y-%m-%d')}")
        sig_p = doc.add_paragraph()
        sig_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig_p.add_run(f"Signed by: {claim_data.get('signature_name')}\n(Beneficiary Signature)")

        doc.save(output_path)
        return True
    except Exception as e:
        logger.error(f"[DocGen] DOCX generation failed: {e}", exc_info=True)
        return False

def generate_html_application(claim_data: dict, photo_url: str | None, output_path: str) -> bool:
    """Generate the official government application HTML by rendering the templates/printable_application.html."""
    try:
        from jinja2 import Template
        
        template_path = Path(__file__).resolve().parent.parent / "templates" / "printable_application.html"
        with open(template_path, "r", encoding="utf-8") as f:
            template_content = f.read()

        html_content = Template(template_content).render(
            **claim_data,
            passport_photo_url=photo_url,
            current_date=datetime.now().strftime("%d %B %Y"),
            prescriptions_attached=bool(claim_data.get("prescriptions")),
            discharge_summary_attached=bool(claim_data.get("discharge_summary")),
            investigation_reports_attached=bool(claim_data.get("investigation_reports")),
            certificates_attached=bool(claim_data.get("certificates")),
            id_proof_attached=bool(claim_data.get("id_proof")),
            ppo_proof_attached=bool(claim_data.get("ppo_proof"))
        )

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        return True
    except Exception as e:
        logger.error(f"[DocGen] HTML generation failed: {e}", exc_info=True)
        return False

def generate_and_upload_application(claim_id: str, claim_data: dict, photo_url: str | None = None) -> dict:
    """
    Main orchestrator for Phase 3 document generation.
    Generates PDF, DOCX, and HTML application documents, uploads them to Supabase, 
    and returns their public URLs.
    """
    results = {}
    temp_photo_path = None

    try:
        # Download photo locally for rendering
        if photo_url:
            temp_photo_path = _download_image(photo_url)

        temp_dir = tempfile.gettempdir()
        
        # 1. Generate PDF
        pdf_filename = f"Application_{claim_id[-8:]}.pdf"
        pdf_temp_path = os.path.join(temp_dir, pdf_filename)
        if generate_pdf_application(claim_data, temp_photo_path, pdf_temp_path):
            public_pdf = upload_to_storage(pdf_temp_path, filename=pdf_filename, bucket_name=Config.SUPABASE_LETTER_BUCKET, folder="applications")
            if public_pdf:
                results["pdf_url"] = public_pdf
            try:
                os.remove(pdf_temp_path)
            except Exception:
                pass

        # 2. Generate DOCX
        docx_filename = f"Application_{claim_id[-8:]}.docx"
        docx_temp_path = os.path.join(temp_dir, docx_filename)
        if generate_docx_application(claim_data, temp_photo_path, docx_temp_path):
            public_docx = upload_to_storage(docx_temp_path, filename=docx_filename, bucket_name=Config.SUPABASE_LETTER_BUCKET, folder="applications")
            if public_docx:
                results["docx_url"] = public_docx
            try:
                os.remove(docx_temp_path)
            except Exception:
                pass

        # 3. Generate HTML
        html_filename = f"Application_{claim_id[-8:]}.html"
        html_temp_path = os.path.join(temp_dir, html_filename)
        if generate_html_application(claim_data, photo_url, html_temp_path):
            public_html = upload_to_storage(html_temp_path, filename=html_filename, bucket_name=Config.SUPABASE_LETTER_BUCKET, folder="applications")
            if public_html:
                results["html_url"] = public_html
            try:
                os.remove(html_temp_path)
            except Exception:
                pass

    except Exception as exc:
        logger.error(f"[DocGen] Application documents generation failed: {exc}", exc_info=True)

    finally:
        # Clean up temp photo
        if temp_photo_path and os.path.exists(temp_photo_path):
            try:
                os.remove(temp_photo_path)
            except Exception:
                pass

    return results
